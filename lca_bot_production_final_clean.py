
import os
import openai
from openai import OpenAI
import requests
from bs4 import BeautifulSoup
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import datetime
import random
from io import BytesIO

# Load API key securely
openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

model_choice = st.sidebar.selectbox("Choose AI Model", ["gpt-3.5-turbo", "gpt-4-turbo"])

def sanitize_input(text):
    return text.strip().replace("<", "").replace(">", "")

def generate_lci_data():
    return pd.DataFrame({
        'Life Cycle Stage': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy Use (MJ)': [random.uniform(80, 120), random.uniform(50, 100), random.uniform(10, 20), random.uniform(15, 30)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(5, 10), random.uniform(8, 12), random.uniform(1, 3), random.uniform(2, 4)],
        'Water Use (L)': [random.uniform(20, 40), random.uniform(10, 30), random.uniform(1, 5), random.uniform(5, 15)]
    })

def load_uploaded_csv(file):
    try:
        return pd.read_csv(file)
    except Exception:
        return None

def scrape_product_data(product):
    query = product.replace(" ", "+") + "+environmental+impact"
    url = f"https://www.google.com/search?q={query}"
    headers = {"User-Agent": "Mozilla/5.0"}
    res = requests.get(url, headers=headers)
    soup = BeautifulSoup(res.text, "html.parser")
    paragraphs = soup.find_all("div", class_="BNeawe s3v9rd AP7Wnd")
    return " ".join([p.get_text() for p in paragraphs[:5]])

def generate_ai_section(prompt, product, model="gpt-3.5-turbo"):
    if not client:
        return f"[Fallback] {prompt} content for {product} would go here."
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a sustainability analyst writing ISO-style LCA reports."},
                {"role": "user", "content": f"Write the '{prompt}' section for a life cycle assessment of a {product}, with citations if possible."}
            ],
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Fallback] Unable to generate content for '{prompt}': {e}"

def create_visuals(df):
    chart_files = []
    for col in df.columns[1:]:
        fig, ax = plt.subplots()
        ax.bar(df['Life Cycle Stage'], df[col], color='green')
        ax.set_title(f"{col} by Life Cycle Stage")
        chart_name = f"{col.replace(' ', '_')}.png"
        fig.savefig(chart_name)
        chart_files.append(chart_name)
        plt.close(fig)
    return chart_files

def create_report(product, df, charts, web_info, ai_sections):
    doc = Document()
    doc.add_heading(f"LCA Report for: {product}", 0)
    doc.add_paragraph(f"Date: {datetime.date.today()}")
    doc.add_paragraph("Confidential — For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for s in ["Executive Summary", "1. Introduction", "2. Goal and Scope", "3. Functional Unit", "4. System Boundary",
              "5. Web Data", "6. Inventory Analysis", "7. LCIA", "8. Interpretation", "9. Limitations",
              "10. Recommendations", "Appendix A: Glossary", "Appendix B: References"]:
        doc.add_paragraph(s)
    doc.add_page_break()

    for section, text in ai_sections.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    doc.add_heading("5. Web Data", level=1)
    doc.add_paragraph(web_info or "No web content found.")
    doc.add_page_break()

    doc.add_heading("6. Inventory Analysis", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_page_break()

    doc.add_heading("7. LCIA", level=1)
    for chart in charts:
        doc.add_paragraph(f"Chart: {chart.split('.')[0]}")
        doc.add_picture(chart, width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading("Appendix A: Glossary", level=1)
    doc.add_paragraph("LCA: Life Cycle Assessment\nGWP: Global Warming Potential\nMJ: Megajoules\nCO2-eq: Carbon dioxide equivalent")
    doc.add_page_break()

    doc.add_heading("Appendix B: References", level=1)
    doc.add_paragraph("1. ISO 14040/44\n2. Ecoinvent\n3. OpenLCA\n4. Public sources scraped from web.")
    file_path = f"LCA_Report_{product.replace(' ', '_')}.docx"
    doc.save(file_path)
    return file_path

st.title("🌿 Production-Grade ISO LCA Bot")

uploaded_file = st.file_uploader("Upload product CSV (optional)", type=["csv"])
product_name = st.text_input("Enter product name (used in the report)", "Electric Toothbrush")
run_btn = st.button("Generate Full Report")

if run_btn:
    product_name = sanitize_input(product_name)
    df = load_uploaded_csv(uploaded_file) if uploaded_file else generate_lci_data()
    charts = create_visuals(df)
    web_data = scrape_product_data(product_name)

    sections = ["Executive Summary", "1. Introduction", "2. Goal and Scope",
                "3. Functional Unit", "4. System Boundary", "8. Interpretation",
                "9. Limitations", "10. Recommendations"]
    ai_content = {sec: generate_ai_section(sec, product_name, model_choice) for sec in sections}

    report_path = create_report(product_name, df, charts, web_data, ai_content)

    with open(report_path, "rb") as f:
        st.download_button("📥 Download DOCX Report", f, file_name=report_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
